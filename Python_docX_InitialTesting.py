from docx import Document

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches

data = {
     "Title" : "Documment Title",
     "Content" : [
         {
             "Chapter": "Chapter 1 Text",
             "Text" : "This is my text1 \nThist is still my text 1",
             "Image": "802.11ac (20)5180.PNG",
             "Table": ["A 123","B321","C231"],
         },
         {
             "Chapter": "Chapter 2 Text",
             "Text" : "This is my text 2 \nThist is still my text 2",
             "Image": "802.11ac (20)5240.PNG",
             "Table": ["A456","B654","C564"],
         },
         {
             "Chapter": "Chapter 2 Text",
             "Text" : "This is my text 2 \nThist is still my text 2",
             "Image": "802.11ac (20)5320.PNG",
             "Table": ["A789","B987","C897"],
         },
     ],
  }

document = Document()

styles = document.styles
#Style paragraph
p = styles.add_style("Paragraph", WD_STYLE_TYPE.PARAGRAPH)
p.font.name = "Arial"
p.font.size = Pt(10)
#Style Heading 2
h2 = styles.add_style("H2", WD_STYLE_TYPE.PARAGRAPH)
h2.base_style = styles ["Heading 2"]
h2.font.name = "Arial"
h2.font.size = Pt(12)
h2.font.color.rgb = RGBColor(79, 129, 189)
h2.font.bold = False
#Style Heading 3
h3 = styles.add_style("H3", WD_STYLE_TYPE.PARAGRAPH)
h3.base_style = styles ["Heading 3"]
h3.font.name = "Arial"
h3.font.size = Pt(14)
h3.font.color.rgb = RGBColor(79, 129, 189)
h3.font.bold = False

Titulo = document.add_heading(data.get("Title"),0)
Titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

for content in data.get("Content"):
    #Add paragraph with h2
    document.add_paragraph(content.get("Chapter"), style="H2")
    #Add paragraph
    document.add_paragraph(content.get("Text"), style="Paragraph")
    #Add paragraph with h3
    paragraph = document.add_paragraph("Image", style="H3")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    my_image = document.add_picture(content.get("Image"), width=Inches(4.25),height=Inches(3.25))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add paragraph with h3



    Tob = document.add_paragraph("Table", style="H3")
    Tob.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = document.add_table(rows=1, cols=3, style="Table Grid")
    table.aligment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Col1"
    hdr_cells[1].text = "Col2"
    hdr_cells[2].text = "Col3"
    row_cells = table.add_row().cells

    for index, element in enumerate(content.get("Table")):
        row_cells[index].text = element



document.save("SaveNewDocxB.docx")