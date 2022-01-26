from docx import Document

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches, Cm


document = Document()

# Define as margens da pagina
sections = document.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
# Definição das margens da pagina

styles = document.styles

Pontificia = styles.add_style("Paragraph", WD_STYLE_TYPE.PARAGRAPH)
Pontificia.font.name = "Arial"
Pontificia.font.size = Pt(14)
Pontificia.font.bold = True

pic = "logo_rodape_novo.png"
pic_2 = "Selo-Inmetro_CRL0075a.png"


table = document.add_table(rows=5, cols=3)
table.allow_autofit = True
row = table.rows[1]
a, b, c = row.cells
AmoreB = a.merge(b)
O = c.merge(AmoreB)

O.height = Cm(0.27)
O.width = Cm(18.18)


row2 = table.rows[2]
a, b, c = row2.cells
AmoreB = a.merge(b)
c.merge(AmoreB)

row3 = table.rows[4]
a, b, c = row3.cells
AmoreB = a.merge(b)
c.merge(AmoreB)



cell = table.rows[0].cells[0]
paragraph = cell.paragraphs[0]
run = paragraph.add_run()

run.add_picture(pic,width = Cm(1.74), height = Cm(2.99))
Celling = table.rows[0].cells

P = Celling[1].add_paragraph("Pontificia Universidade Catolica do Rio Grande do Sul\n", style = "Paragraph")
P.add_run("LABELO - Laboratorios Especializados em Eletroeletronica\n").font.size = Pt(12)
P.add_run("Calibracao e Ensaios").font.size = Pt(11)

Celling[1].width = Cm(13.46)
Celling[2].width = Cm(1.79)

table.alignment = WD_TABLE_ALIGNMENT.CENTER
P.alignment = WD_TABLE_ALIGNMENT.CENTER

Cell2 = table.rows[0].cells[2]


paragraph_2 = Cell2.paragraphs[0]

paragraph_2.alignment = WD_TABLE_ALIGNMENT.RIGHT

Running = paragraph_2.add_run()
Celling[0].width = Cm(1.77)
Running.add_picture(pic_2,width = Cm(1.79), height = Cm(2.98))
Running.alignment = WD_TABLE_ALIGNMENT.CENTER
document.save("LABELO_Test.docx")